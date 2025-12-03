"""
ZMtech Advanced Stock Analysis Platform
Combines technical analysis with AI-powered equity research
Supports both Streamlit secrets and .env file for API key management

Version: 2.0 (with lazy import for Streamlit Cloud compatibility)
"""

import streamlit as st
import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import yfinance as yf
from datetime import datetime, timedelta
import numpy as np
from statsmodels.tsa.statespace.sarimax import SARIMAX
import warnings
import base64
import os
import re
import io
from pathlib import Path

# Import for report generation
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Lazy import function for FullStockAnalyzer to handle deployment issues
# This function is called only when needed, avoiding import-time errors
def get_full_stock_analyzer():
    """Lazy import of FullStockAnalyzer with multiple fallback methods for Streamlit Cloud compatibility"""
    import sys
    import importlib.util
    import os
    
    # Method 1: Standard import (works locally)
    try:
        from full_analysis import FullStockAnalyzer
        return FullStockAnalyzer
    except (ImportError, KeyError, ModuleNotFoundError, AttributeError) as e1:
        pass  # Continue to next method
    
    # Method 2: Direct file import using importlib (for Streamlit Cloud)
    try:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        # Try multiple possible locations
        possible_paths = [
            os.path.join(current_dir, 'full_analysis.py'),
            os.path.join(os.path.dirname(current_dir), 'full_analysis.py'),
            os.path.join(os.getcwd(), 'full_analysis.py'),
        ]
        
        for full_analysis_path in possible_paths:
            if os.path.exists(full_analysis_path):
                spec = importlib.util.spec_from_file_location(
                    "full_analysis_module", 
                    full_analysis_path
                )
                if spec and spec.loader:
                    # Use a unique module name to avoid conflicts
                    module_name = f"full_analysis_{id(spec)}"
                    full_analysis_module = importlib.util.module_from_spec(spec)
                    sys.modules[module_name] = full_analysis_module
                    spec.loader.exec_module(full_analysis_module)
                    if hasattr(full_analysis_module, 'FullStockAnalyzer'):
                        return full_analysis_module.FullStockAnalyzer
    except Exception as e2:
        pass  # Continue to next method
    
    # Method 3: Add directories to path and retry
    try:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        parent_dir = os.path.dirname(current_dir)
        cwd = os.getcwd()
        
        for dir_path in [current_dir, parent_dir, cwd]:
            if dir_path and dir_path not in sys.path:
                sys.path.insert(0, dir_path)
        
        from full_analysis import FullStockAnalyzer
        return FullStockAnalyzer
    except Exception as e3:
        pass
    
    # If all methods fail, raise a comprehensive error
    current_dir = os.path.dirname(os.path.abspath(__file__))
    raise ImportError(
        f"Failed to import FullStockAnalyzer from full_analysis.py. "
        f"Tried: standard import, importlib, and path manipulation. "
        f"Current dir: {os.getcwd()}, "
        f"Script dir: {current_dir}, "
        f"Python path: {sys.path[:3]}"
    )

# Try to import optional dependencies
try:
    from curl_cffi import requests
    CURL_CFFI_AVAILABLE = True
except ImportError:
    CURL_CFFI_AVAILABLE = False

try:
    from dotenv import load_dotenv
    load_dotenv()
    DOTENV_AVAILABLE = True
except ImportError:
    DOTENV_AVAILABLE = False

warnings.filterwarnings('ignore')

# Page configuration
st.set_page_config(
    page_title="ZMtech Stock Analysis Platform",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ═══════════════════════════════════════════════════════════════
# API KEY MANAGEMENT
# Priority: Streamlit secrets → .env file → Environment variable
# ═══════════════════════════════════════════════════════════════

def get_google_api_key():
    """
    Get Google AI API key from multiple sources in priority order:
    1. Streamlit secrets (.streamlit/secrets.toml)
    2. Environment variables (.env file or system)
    3. Return None if not found
    """
    # Try Streamlit secrets first
    try:
        if "GOOGLE_API_KEY" in st.secrets:
            return st.secrets["GOOGLE_API_KEY"]
    except:
        pass
    
    # Try environment variable (from .env or system)
    api_key = os.getenv("GOOGLE_API_KEY")
    if api_key:
        return api_key
    
    # Not found
    return None

# ═══════════════════════════════════════════════════════════════
# UTILITY FUNCTIONS
# ═══════════════════════════════════════════════════════════════

@st.cache_resource
def get_yfinance_session():
    """Create a curl_cffi session with Chrome impersonation for yfinance"""
    if CURL_CFFI_AVAILABLE:
        try:
            session = requests.Session(impersonate="chrome")
            if 'curl_cffi_working' not in st.session_state:
                st.session_state.curl_cffi_working = True
            return session
        except Exception as e:
            if 'curl_cffi_warning' not in st.session_state:
                st.warning(f"Failed to create curl_cffi session: {e}. Using default requests.")
                st.session_state.curl_cffi_warning = True
            return None
    return None

def add_bg_from_local():
    """Add custom CNBC-style finance theme with background image"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    image_path = os.path.join(current_dir, "background_v2.png")
    
    if os.path.exists(image_path):
        try:
            with open(image_path, "rb") as image_file:
                encoded_string = base64.b64encode(image_file.read()).decode()
                
            st.markdown(
                f"""
                <style>
                /* CNBC-style dark finance theme */
                .stApp {{
                    background: #000000;
                    background-image: linear-gradient(rgba(0, 0, 0, 0.85), rgba(0, 0, 0, 0.85)), url(data:image/png;base64,{encoded_string});
                    background-size: cover;
                    background-position: center;
                    background-repeat: no-repeat;
                    background-attachment: fixed;
                }}
                
                /* Main content containers */
                .main .block-container {{
                    background-color: rgba(26, 26, 26, 0.95);
                    padding: 2rem;
                    border-radius: 8px;
                }}
                
                .element-container, .stMarkdown, .stDataFrame {{
                    background-color: rgba(26, 26, 26, 0.9);
                    padding: 1rem;
                    border-radius: 8px;
                    border: 1px solid rgba(255, 255, 255, 0.1);
                }}
                
                /* Text colors - CNBC style */
                h1, h2, h3, h4, h5, h6 {{
                    color: #ffffff !important;
                    font-family: 'Helvetica Neue', Arial, sans-serif;
                    font-weight: 600;
                }}
                
                p, label, div, span {{
                    color: #e0e0e0 !important;
                }}
                
                /* Sidebar styling */
                [data-testid="stSidebar"] {{
                    background-color: #1a1a1a;
                }}
                
                /* Buttons - CNBC blue */
                .stButton>button {{
                    background-color: #0066cc;
                    color: #ffffff;
                    font-weight: 600;
                    border-radius: 4px;
                    border: none;
                    padding: 0.5rem 1.5rem;
                    transition: all 0.2s ease;
                }}
                
                .stButton>button:hover {{
                    background-color: #0052a3;
                    transform: translateY(-1px);
                    box-shadow: 0 4px 8px rgba(0, 102, 204, 0.3);
                }}
                
                /* Input fields */
                .stTextInput>div>div>input {{
                    background-color: rgba(26, 26, 26, 0.8);
                    color: #ffffff;
                    border: 1px solid rgba(255, 255, 255, 0.2);
                    border-radius: 4px;
                }}
                
                .stSelectbox>div>div>select {{
                    background-color: rgba(26, 26, 26, 0.8);
                    color: #ffffff;
                    border: 1px solid rgba(255, 255, 255, 0.2);
                }}
                
                /* Metrics - CNBC style */
                [data-testid="stMetricValue"] {{
                    color: #ffffff !important;
                    font-weight: 700;
                }}
                
                [data-testid="stMetricLabel"] {{
                    color: #b0b0b0 !important;
                }}
                
                /* Data tables */
                .stDataFrame {{
                    background-color: rgba(26, 26, 26, 0.9);
                }}
                
                table {{
                    color: #ffffff;
                }}
                
                /* Positive/negative colors */
                .positive {{
                    color: #00ff00 !important;
                }}
                
                .negative {{
                    color: #ff0000 !important;
                }}
                
                /* Streamlit widgets */
                .stSlider {{
                    color: #ffffff;
                }}
                
                /* Markdown text */
                .stMarkdown {{
                    color: #e0e0e0;
                }}
                
                /* Info boxes */
                .stInfo {{
                    background-color: rgba(0, 102, 204, 0.2);
                    border-left: 4px solid #0066cc;
                }}
                
                .stSuccess {{
                    background-color: rgba(0, 255, 0, 0.1);
                    border-left: 4px solid #00ff00;
                }}
                
                .stError {{
                    background-color: rgba(255, 0, 0, 0.1);
                    border-left: 4px solid #ff0000;
                }}
                
                .stWarning {{
                    background-color: rgba(255, 193, 7, 0.1);
                    border-left: 4px solid #ffc107;
                }}
                </style>
                """,
                unsafe_allow_html=True
            )
        except Exception:
            pass

def format_report_text(report: str) -> str:
    """Format report text to highlight section headers"""
    formatted = report
    
    # Remove excessive newlines (more than 2)
    formatted = re.sub(r'\n{3,}', '\n\n', formatted)
    
    # Format numbered headers (e.g., ## 1. EXECUTIVE SUMMARY)
    # Use div with margins for better spacing
    formatted = re.sub(
        r'(##\s*\d+\.\s*[A-Z\s&]+)',
        r'<div style="color: #0066cc; font-size: 18px; font-weight: 700; margin-top: 25px; margin-bottom: 10px; border-bottom: 1px solid rgba(0, 102, 204, 0.3); padding-bottom: 5px;">\1</div>',
        formatted,
        flags=re.IGNORECASE
    )
    
    # Format main title
    formatted = re.sub(
        r'(#\s*EQUITY RESEARCH NOTE:.*)',
        r'<h1 style="color: #ffffff; font-size: 24px; border-bottom: 2px solid #0066cc; padding-bottom: 10px; margin-bottom: 20px;">\1</h1>',
        formatted,
        flags=re.IGNORECASE
    )
    
    # Bold keys (e.g., **Rating:**)
    formatted = re.sub(
        r'(\*\*[^*]+\*\*:)',
        r'<span style="color: #e0e0e0; font-weight: 700;">\1</span>',
        formatted
    )
    
    return formatted

def create_word_report(report_text, ticker):
    """Generate a formatted Word document from the report text"""
    doc = Document()
    
    # Title
    title = doc.add_heading(f'Equity Research Report: {ticker}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('-------------------------------------------------------------------')
    
    # Process text line by line for basic formatting
    lines = report_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check for headers
        if line.startswith('# '):
            # Main Title
            p = doc.add_heading(line.replace('# ', ''), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            # Section Headers
            p = doc.add_heading(line.replace('## ', ''), level=1)
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0, 102, 204)  # CNBC Blue
        elif line.startswith('**') and line.endswith('**'):
            # Bold lines
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', ''))
            run.bold = True
        elif line.startswith('TICKER:') or line.startswith('CURRENT PRICE:'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        elif 'BUY' in line or 'STRONG BUY' in line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 153, 51)  # Green
        elif 'SELL' in line or 'STRONG SELL' in line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.color.rgb = RGBColor(204, 0, 0)  # Red
        else:
            # Handle inline bolding like **Text**
            if '**' in line:
                parts = line.split('**')
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1: # Odd parts are between ** **
                        run.bold = True
            else:
                doc.add_paragraph(line)
            
    # Footer
    doc.add_paragraph('-------------------------------------------------------------------')
    footer = doc.add_paragraph('Generated by ZMtech Advanced Stock Analysis Platform')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(0, 10, 'ZMtech Equity Research', 0, 1, 'C')
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def create_pdf_report(report_text, ticker):
    """Generate a formatted PDF document from the report text"""
    pdf = PDF()
    pdf.add_page()
    
    # Title
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, f"Equity Research Report: {ticker}", 0, 1, "C")
    pdf.set_font("Arial", "I", 10)
    pdf.cell(0, 10, f"Generated on: {datetime.now().strftime('%B %d, %Y')}", 0, 1, "C")
    pdf.ln(5)
    
    # Content
    pdf.set_font("Courier", size=10) # Use monospaced font for alignment
    
    # Clean up text for PDF (remove some special chars that might cause issues)
    clean_text = report_text.encode('latin-1', 'replace').decode('latin-1')
    
    pdf.multi_cell(0, 5, clean_text)
    
    # Output to bytes
    return pdf.output(dest='S').encode('latin-1')

# ═══════════════════════════════════════════════════════════════
# TECHNICAL ANALYSIS FUNCTIONS
# ═══════════════════════════════════════════════════════════════

def calculate_ema(data, period):
    return data['Close'].ewm(span=period, adjust=False).mean()

def calculate_vwap(df):
    df = df.copy()
    typical_price = (df['High'] + df['Low'] + df['Close']) / 3
    volume_price = typical_price * df['Volume']
    cumulative_volume = df['Volume'].cumsum()
    cumulative_volume_price = volume_price.cumsum()
    return cumulative_volume_price / cumulative_volume

def calculate_macd(data):
    exp1 = data['Close'].ewm(span=12, adjust=False).mean()
    exp2 = data['Close'].ewm(span=26, adjust=False).mean()
    macd = exp1 - exp2
    signal = macd.ewm(span=9, adjust=False).mean()
    return macd, signal

def calculate_rsi(data, periods=14):
    delta = data['Close'].diff()
    gain = (delta.where(delta > 0, 0)).rolling(window=periods).mean()
    loss = (-delta.where(delta < 0, 0)).rolling(window=periods).mean()
    rs = gain / loss
    return 100 - (100 / (1 + rs))

def generate_signals(df):
    """Generate buy and sell signals"""
    df = df.copy()
    
    df['Buy_Signal'] = ((df['MACD'] > df['Signal']) & 
                        (df['MACD'].shift(1) <= df['Signal'].shift(1)) & 
                        (df['RSI'] < 70)).astype(int)
    
    df['Sell_Signal'] = ((df['MACD'] < df['Signal']) & 
                         (df['MACD'].shift(1) >= df['Signal'].shift(1)) & 
                         (df['RSI'] > 30)).astype(int)
    
    return df

def calculate_heikin_ashi(df):
    """Calculate Heikin-Ashi candlestick data"""
    ha_close = (df['Open'] + df['High'] + df['Low'] + df['Close']) / 4
    
    ha_open = pd.Series([(df['Open'].iloc[0] + df['Close'].iloc[0]) / 2], index=[df.index[0]])
    
    for i in range(1, len(df)):
        next_value = pd.Series(
            [(ha_open.iloc[-1] + ha_close.iloc[i-1]) / 2],
            index=[df.index[i]]
        )
        ha_open = pd.concat([ha_open, next_value])
    
    ha_high = pd.concat([df['High'], ha_open, ha_close], axis=1).max(axis=1)
    ha_low = pd.concat([df['Low'], ha_open, ha_close], axis=1).min(axis=1)
    
    return pd.DataFrame({
        'HA_Open': ha_open,
        'HA_High': ha_high,
        'HA_Low': ha_low,
        'HA_Close': ha_close
    }, index=df.index)

def forecast_sarima(data, periods=30):
    """Generate SARIMA forecast"""
    try:
        model = SARIMAX(data['Close'], order=(1, 1, 1), seasonal_order=(1, 1, 1, 12))
        results = model.fit(disp=False)
        forecast = results.forecast(steps=periods)
        return forecast
    except:
        return None

@st.cache_data(ttl=300)
def download_stock_data_cached(ticker, start_date, end_date):
    """Cached wrapper for stock data download"""
    session = get_yfinance_session()
    
    try:
        if session:
            ticker_obj = yf.Ticker(ticker, session=session)
        else:
            ticker_obj = yf.Ticker(ticker)
        data = ticker_obj.history(start=start_date, end=end_date)
        
        if data is not None and not data.empty:
            if isinstance(data.columns, pd.MultiIndex):
                data.columns = data.columns.get_level_values(0)
            return data
    except:
        pass
    
    return None

# ═══════════════════════════════════════════════════════════════
# CUSTOM CSS STYLING
# ═══════════════════════════════════════════════════════════════

st.markdown("""
<style>
    /* CNBC-style finance theme - additional styling */
    .main {
        background: #000000;
    }
    
    .metric-card {
        background: rgba(26, 26, 26, 0.9);
        backdrop-filter: blur(10px);
        border-radius: 8px;
        padding: 15px;
        margin: 10px 0;
        border: 1px solid rgba(255,255,255,0.1);
        color: #e0e0e0;
    }
    
    .report-container {
        background: rgba(26, 26, 26, 0.95);
        color: #ffffff;
        padding: 30px;
        border-radius: 8px;
        font-family: 'Helvetica Neue', Arial, sans-serif;
        box-shadow: 0 4px 12px rgba(0,0,0,0.5);
        border: 1px solid rgba(255,255,255,0.1);
        margin: 20px 0;
    }
    
    /* Additional CNBC styling */
    .stTabs [data-baseweb="tab-list"] {
        background-color: rgba(26, 26, 26, 0.9);
    }
    
    .stTabs [data-baseweb="tab"] {
        color: #e0e0e0;
    }
    
    .stTabs [aria-selected="true"] {
        color: #0066cc;
        border-bottom: 2px solid #0066cc;
    }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════
# MAIN APPLICATION
# ═══════════════════════════════════════════════════════════════

# Header - CNBC style
st.markdown("""
    <h1 style='text-align: center; font-size: 3em; margin-bottom: 0; color: #ffffff; font-weight: 700; font-family: "Helvetica Neue", Arial, sans-serif;'>
        📊 ZMtech Stock Analysis Platform
    </h1>
    <p style='text-align: center; color: #b0b0b0; font-size: 1.2em; margin-top: 0; font-weight: 400;'>
        Advanced Technical Analysis + AI-Powered Equity Research
    </p>
""", unsafe_allow_html=True)

st.markdown("---")

# Sidebar configuration
with st.sidebar:
    st.markdown("### ⚙️ Configuration")
    
    # Get API key
    google_api_key = get_google_api_key()
    
    # Show API key status and allow override
    if google_api_key:
        st.success("✅ Google AI API key found")
        use_custom_key = st.checkbox("Use different API key")
        if use_custom_key:
            google_api_key = st.text_input(
                "Custom Google AI API Key",
                type="password",
                help="Override the configured API key"
            )
    else:
        st.warning("⚠️ Google AI API key not configured")
        google_api_key = st.text_input(
            "Google AI API Key (Required for AI Analysis)",
            type="password",
            help="Get your free API key from https://makersuite.google.com/app/apikey"
        )
        
        if not google_api_key:
            st.info("""
            **Setup API Key:**
            1. Get free API key from [Google AI Studio](https://makersuite.google.com/app/apikey)
            2. Add to `.streamlit/secrets.toml`:
               ```
               GOOGLE_API_KEY = "your-key"
               ```
            3. Or create `.env` file:
               ```
               GOOGLE_API_KEY=your-key
               ```
            """)
    
    st.markdown("---")
    
    # Analysis mode selection
    st.markdown("### 📊 Analysis Mode")
    analysis_mode = st.radio(
        "Select Mode",
        ["Technical Analysis", "AI-Powered Research", "Combined Analysis"],
        help="Choose your analysis approach"
    )
    
    st.markdown("---")
    
    # Stock selection
    st.markdown("### 📈 Stock Selection")
    ticker = st.text_input(
        "Stock Ticker",
        value="AAPL",
        help="Enter stock ticker symbol"
    ).upper().strip()
    
    # Date range
    today = datetime.today().date()
    min_date = today - timedelta(days=365*10)
    
    start_date = st.date_input(
        "Start Date",
        value=today - timedelta(days=365),
        min_value=min_date,
        max_value=today
    )
    
    end_date = st.date_input(
        "End Date",
        value=today,
        min_value=start_date,
        max_value=today
    )
    
    # Chart options (for technical analysis)
    if analysis_mode in ["Technical Analysis", "Combined Analysis"]:
        st.markdown("---")
        st.markdown("### 📉 Chart Options")
        
        chart_type = st.selectbox(
            "Chart Type",
            ["Regular Candlestick", "Heikin-Ashi"]
        )
        
        show_forecast = st.checkbox("Show SARIMA Forecast")
        if show_forecast:
            forecast_periods = st.slider("Forecast Periods", 5, 60, 30)
    
    st.markdown("---")
    
    # Analysis button
    analyze_button = st.button("🚀 Generate Analysis", use_container_width=True)

# ═══════════════════════════════════════════════════════════════
# MAIN CONTENT AREA
# ═══════════════════════════════════════════════════════════════

if not ticker:
    st.error("⚠️ Please enter a stock ticker symbol")
    st.stop()

if analyze_button:
    with st.spinner(f'🔍 Analyzing {ticker}...'):
        # Fetch stock data
        try:
            download_end = min(end_date, today) + timedelta(days=1)
            data = download_stock_data_cached(ticker, start_date, download_end)
            
            if data is None or data.empty:
                st.error(f"❌ No data found for {ticker}. Please check the ticker symbol.")
                st.stop()
            
            # Filter future dates
            if not data.empty:
                last_valid_date = data.index.max().date()
                if last_valid_date > today:
                    data = data[data.index.date <= today]
            
            if data.empty or len(data) < 26:
                st.error("❌ Not enough data for analysis. Please select a longer date range.")
                st.stop()
            
            # Calculate technical indicators
            data = data.copy()
            data['EMA9'] = calculate_ema(data, 9)
            data['EMA20'] = calculate_ema(data, 20)
            data['EMA50'] = calculate_ema(data, 50)
            data['EMA100'] = calculate_ema(data, 100)
            data['EMA200'] = calculate_ema(data, 200)
            data['VWAP'] = calculate_vwap(data)
            macd, signal = calculate_macd(data)
            data['MACD'] = macd
            data['Signal'] = signal
            data['RSI'] = calculate_rsi(data)
            
            # Heikin-Ashi
            ha_df = calculate_heikin_ashi(data)
            data = pd.concat([data, ha_df], axis=1)
            
            # Signals
            data = generate_signals(data)
            
            # Calculate current metrics
            current_price = data['Close'].iloc[-1]
            price_change = ((current_price - data['Close'].iloc[-2]) / data['Close'].iloc[-2] * 100) if len(data) > 1 else 0
            current_rsi = data['RSI'].iloc[-1]
            current_volume = data['Volume'].iloc[-1]
            avg_volume = data['Volume'].tail(20).mean()
            
            # Display quick metrics
            st.markdown("### 📊 Quick Metrics")
            col1, col2, col3, col4, col5 = st.columns(5)
            
            with col1:
                st.metric("Current Price", f"${current_price:.2f}", f"{price_change:.2f}%")
            
            with col2:
                rsi_status = "🔴 Overbought" if current_rsi > 70 else "🟢 Oversold" if current_rsi < 30 else "🟡 Neutral"
                st.metric("RSI (14)", f"{current_rsi:.1f}", rsi_status)
            
            with col3:
                ema_trend = "🟢 Bullish" if data['EMA20'].iloc[-1] > data['EMA50'].iloc[-1] else "🔴 Bearish"
                st.metric("Trend (EMA)", ema_trend)
            
            with col4:
                vol_trend = "🟢 High" if current_volume > avg_volume * 1.5 else "🟡 Normal" if current_volume > avg_volume * 0.7 else "🔴 Low"
                st.metric("Volume", vol_trend, f"{current_volume/avg_volume:.2f}x avg")
            
            with col5:
                macd_signal_text = "🟢 Bullish" if data['MACD'].iloc[-1] > data['Signal'].iloc[-1] else "🔴 Bearish"
                st.metric("MACD Signal", macd_signal_text)
            
            st.markdown("---")
            
            # ═══════════════════════════════════════════════════════════
            # TECHNICAL ANALYSIS MODE
            # ═══════════════════════════════════════════════════════════
            
            if analysis_mode in ["Technical Analysis", "Combined Analysis"]:
                st.markdown("### 📈 Technical Analysis Chart")
                
                # Create chart
                fig = make_subplots(
                    rows=4, cols=1,
                    shared_xaxes=True,
                    vertical_spacing=0.05,
                    row_heights=[0.5, 0.15, 0.15, 0.15],
                    subplot_titles=('Price & Indicators', 'Volume', 'MACD', 'RSI')
                )
                
                # Candlestick
                if chart_type == "Regular Candlestick":
                    candlestick_data = dict(
                        x=data.index, open=data['Open'], high=data['High'],
                        low=data['Low'], close=data['Close'], name='Price'
                    )
                else:
                    candlestick_data = dict(
                        x=data.index, open=data['HA_Open'], high=data['HA_High'],
                        low=data['HA_Low'], close=data['HA_Close'], name='Heikin-Ashi'
                    )
                
                fig.add_trace(
                    go.Candlestick(**candlestick_data, increasing_line_color='green', decreasing_line_color='red'),
                    row=1, col=1
                )
                
                # EMAs
                ema_colors = {'EMA9': '#ffeb3b', 'EMA20': '#00bcd4', 'EMA50': '#ff9800', 
                             'EMA100': '#9c27b0', 'EMA200': '#f44336'}
                for ema, color in ema_colors.items():
                    if ema in data.columns:
                        fig.add_trace(
                            go.Scatter(x=data.index, y=data[ema], name=ema, line=dict(color=color, width=1.5)),
                            row=1, col=1
                        )
                
                # VWAP
                fig.add_trace(
                    go.Scatter(x=data.index, y=data['VWAP'], name='VWAP', line=dict(color='purple', width=1.5)),
                    row=1, col=1
                )
                
                # SARIMA Forecast
                if show_forecast:
                    forecast = forecast_sarima(data, periods=forecast_periods)
                    if forecast is not None:
                        fig.add_trace(
                            go.Scatter(
                                x=pd.date_range(start=data.index[-1], periods=len(forecast)+1)[1:],
                                y=forecast, name='SARIMA Forecast',
                                line=dict(color='orange', dash='dash')
                            ),
                            row=1, col=1
                        )
                
                # Buy/Sell Signals
                buy_mask = data['Buy_Signal'] == 1
                if buy_mask.any():
                    buy_signals = data[buy_mask]
                    fig.add_trace(
                        go.Scatter(
                            x=buy_signals.index, y=buy_signals['Low'] * 0.99,
                            mode='markers+text', marker=dict(symbol='triangle-up', size=15, color='green'),
                            text='BUY', textposition='bottom center', name='Buy Signal'
                        ),
                        row=1, col=1
                    )
                
                sell_mask = data['Sell_Signal'] == 1
                if sell_mask.any():
                    sell_signals = data[sell_mask]
                    fig.add_trace(
                        go.Scatter(
                            x=sell_signals.index, y=sell_signals['High'] * 1.01,
                            mode='markers+text', marker=dict(symbol='triangle-down', size=15, color='red'),
                            text='SELL', textposition='top center', name='Sell Signal'
                        ),
                        row=1, col=1
                    )
                
                # Volume
                colors_vol = ['green' if c > o else 'red' for c, o in zip(data['Close'], data['Open'])]
                fig.add_trace(
                    go.Bar(x=data.index, y=data['Volume'], name='Volume', marker_color=colors_vol, opacity=0.7),
                    row=2, col=1
                )
                
                # MACD
                fig.add_trace(
                    go.Scatter(x=data.index, y=data['MACD'], name='MACD', line=dict(color='blue', width=2)),
                    row=3, col=1
                )
                fig.add_trace(
                    go.Scatter(x=data.index, y=data['Signal'], name='Signal', line=dict(color='orange', width=2)),
                    row=3, col=1
                )
                macd_hist_colors = ['green' if v > 0 else 'red' for v in (data['MACD'] - data['Signal'])]
                fig.add_trace(
                    go.Bar(x=data.index, y=data['MACD'] - data['Signal'], name='MACD Histogram', 
                          marker_color=macd_hist_colors, opacity=0.5),
                    row=3, col=1
                )
                
                # RSI
                fig.add_trace(
                    go.Scatter(x=data.index, y=data['RSI'], name='RSI', line=dict(color='#9c27b0', width=2)),
                    row=4, col=1
                )
                fig.add_hline(y=70, line_dash="dash", line_color="red", opacity=0.5, row=4, col=1)
                fig.add_hline(y=30, line_dash="dash", line_color="green", opacity=0.5, row=4, col=1)
                fig.add_hline(y=50, line_dash="dot", line_color="gray", opacity=0.3, row=4, col=1)
                
                # Layout
                fig.update_layout(
                    title=f"{ticker} Technical Analysis",
                    template='plotly_dark',
                    height=1000,
                    showlegend=True,
                    hovermode='x unified',
                    xaxis_rangeslider_visible=False
                )
                
                st.plotly_chart(fig, width='stretch')
            
            # ═══════════════════════════════════════════════════════════
            # AI-POWERED RESEARCH MODE
            # ═══════════════════════════════════════════════════════════
            
            if analysis_mode in ["AI-Powered Research", "Combined Analysis"]:
                st.markdown("---")
                st.markdown("### 🤖 AI-Powered Equity Research Report")
                
                if not google_api_key:
                    st.error("❌ Google AI API key is required for AI analysis")
                    st.info("Please enter your API key in the sidebar or configure it in secrets.toml")
                else:
                    try:
                        with st.status("Generating AI analysis...", expanded=True) as status:
                            st.write("🤖 Initializing AI analyzer...")
                            try:
                                FullStockAnalyzer = get_full_stock_analyzer()
                            except Exception as import_error:
                                st.error(f"❌ Failed to import FullStockAnalyzer: {import_error}")
                                st.error("Please ensure full_analysis.py is in the same directory as main.py")
                                status.update(label="❌ Import Error", state="error", expanded=False)
                                st.stop()
                            
                            session = get_yfinance_session()
                            analyzer = FullStockAnalyzer(api_key=google_api_key, session=session)
                            
                            st.write("📊 Fetching comprehensive stock data...")
                            # Calculate period for AI analysis
                            days_diff = (end_date - start_date).days
                            if days_diff <= 60:
                                period = "1mo"
                            elif days_diff <= 180:
                                period = "3mo"
                            elif days_diff <= 365:
                                period = "6mo"
                            elif days_diff <= 730:
                                period = "1y"
                            else:
                                period = "2y"
                            
                            stock_data = analyzer.fetch_stock_data(ticker, period)
                            st.write("✅ Data retrieved")
                            
                            st.write("🧠 Generating AI-powered analysis...")
                            report = analyzer.generate_analysis_report(stock_data)
                            st.write("✅ Analysis complete")
                            
                            status.update(label="✅ AI Analysis Complete!", state="complete", expanded=False)
                        
                        # Format and display report
                        formatted_report = format_report_text(report)
                        st.markdown(f"""
                        <div class="report-container">
                            <div style="color: #e0e0e0; font-size: 15px; line-height: 1.6; white-space: pre-wrap; font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;">
{formatted_report}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Download buttons
                        st.markdown("### 📥 Download Report")
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            st.download_button(
                                label="📄 Download as Text",
                                data=report,
                                file_name=f"{ticker}_analysis_{datetime.now().strftime('%Y%m%d')}.txt",
                                mime="text/plain",
                                use_container_width=True
                            )
                            
                        with col2:
                            if WORD_AVAILABLE:
                                try:
                                    word_data = create_word_report(report, ticker)
                                    st.download_button(
                                        label="📝 Download as Word",
                                        data=word_data,
                                        file_name=f"{ticker}_analysis_{datetime.now().strftime('%Y%m%d')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        use_container_width=True
                                    )
                                except Exception as e:
                                    st.error(f"Word generation failed: {e}")
                            else:
                                st.warning("Word export unavailable")
                                
                        with col3:
                            if PDF_AVAILABLE:
                                try:
                                    pdf_data = create_pdf_report(report, ticker)
                                    st.download_button(
                                        label="📕 Download as PDF",
                                        data=pdf_data,
                                        file_name=f"{ticker}_analysis_{datetime.now().strftime('%Y%m%d')}.pdf",
                                        mime="application/pdf",
                                        use_container_width=True
                                    )
                                except Exception as e:
                                    st.error(f"PDF generation failed: {e}")
                            else:
                                st.warning("PDF export unavailable")
                        
                        # Support/Resistance levels
                        st.markdown("### 🎯 Key Price Levels")
                        col1, col2 = st.columns(2)
                        
                        sr = stock_data['support_resistance']
                        with col1:
                            st.markdown("#### 🔴 Resistance Levels")
                            st.markdown(f"""
                            - **R3 (Strong):** ${sr['resistance_3']:.2f}
                            - **R2 (Moderate):** ${sr['resistance_2']:.2f}
                            - **R1 (Immediate):** ${sr['resistance_1']:.2f}
                            """)
                        
                        with col2:
                            st.markdown("#### 🟢 Support Levels")
                            st.markdown(f"""
                            - **S1 (Immediate):** ${sr['support_1']:.2f}
                            - **S2 (Moderate):** ${sr['support_2']:.2f}
                            - **S3 (Strong):** ${sr['support_3']:.2f}
                            """)
                        
                        st.info(f"**Pivot Point:** ${sr['pivot']:.2f}")
                        
                    except Exception as e:
                        st.error(f"❌ Error generating AI analysis: {str(e)}")
                        st.info("Please check your API key and try again.")
            
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
            st.info("Please try again or check your inputs.")

else:
    # Welcome screen
    st.markdown("""
    <div style='text-align: center; padding: 50px;'>
        <h2>👈 Configure your analysis in the sidebar and click "Generate Analysis"</h2>
        <p style='font-size: 1.1em; color: #888; margin-top: 20px;'>
            Choose from three powerful analysis modes:
        </p>
        <div style='display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 20px; margin-top: 30px;'>
            <div class='metric-card'>
                <h3>📈 Technical Analysis</h3>
                <p>Advanced charting with 15+ indicators, buy/sell signals, and SARIMA forecasting</p>
            </div>
            <div class='metric-card'>
                <h3>🤖 AI-Powered Research</h3>
                <p>Professional equity research reports with Buy/Hold/Sell recommendations</p>
            </div>
            <div class='metric-card'>
                <h3>🎯 Combined Analysis</h3>
                <p>Get both technical charts and AI insights in one comprehensive view</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666; padding: 20px;'>
    <p>⚡ Powered by Google Generative AI & Yahoo Finance | ZMtech Analysis Platform</p>
    <p style='font-size: 0.8em;'>Disclaimer: For informational purposes only. Not investment advice. Consult a financial advisor.</p>
</div>
""", unsafe_allow_html=True)
