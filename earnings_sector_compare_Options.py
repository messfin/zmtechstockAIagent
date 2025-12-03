import yfinance as yf
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import tkinter as tk
from tkinter import ttk, messagebox
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
import warnings
warnings.filterwarnings('ignore')

class OptionsAnalyzer:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Options Analysis")
        
        # Create main frames
        self.input_frame = ttk.Frame(self.root, padding="10")
        self.input_frame.grid(row=0, column=0, sticky="nsew")
        
        self.analysis_frame = ttk.Frame(self.root, padding="10")
        self.analysis_frame.grid(row=1, column=0, sticky="nsew")
        
        self.chart_frame = ttk.Frame(self.root, padding="10")
        self.chart_frame.grid(row=2, column=0, sticky="nsew")
        
        self.create_widgets()
        
    def create_widgets(self):
        # Input Frame
        ttk.Label(self.input_frame, text="Ticker:").grid(row=0, column=0)
        self.ticker_entry = ttk.Entry(self.input_frame)
        self.ticker_entry.grid(row=0, column=1)
        
        ttk.Label(self.input_frame, text="Analysis Type:").grid(row=1, column=0)
        self.analysis_type = ttk.Combobox(self.input_frame, values=[
            "Historical IV Analysis",
            "Options Chain Analysis",
            "Strategy Analysis"
        ], state='readonly')
        self.analysis_type.grid(row=1, column=1)
        self.analysis_type.set("Historical IV Analysis")
        
        ttk.Button(self.input_frame, text="Analyze", 
                  command=self.run_analysis).grid(row=2, column=0, columnspan=2)
        
        # Export buttons
        ttk.Button(self.input_frame, text="Export Data", 
                  command=self.export_data).grid(row=3, column=0)
        ttk.Button(self.input_frame, text="Export Report", 
                  command=self.export_report).grid(row=3, column=1)

    def run_analysis(self):
        ticker = self.ticker_entry.get().strip().upper()
        if not ticker:
            messagebox.showerror("Error", "Please enter a ticker symbol")
            return
            
        try:
            stock = yf.Ticker(ticker)
            
            # Get historical data and calculate IV
            hist_data = stock.history(period='1y')
            hist_data['Returns'] = hist_data['Close'].pct_change()
            hist_data['Historical_IV'] = (
                hist_data['Returns'].rolling(window=20).std() * 
                np.sqrt(252) * 100
            )
            
            # Clear previous analysis
            for widget in self.analysis_frame.winfo_children():
                widget.destroy()
            
            analysis_type = self.analysis_type.get()
            
            if analysis_type == "Historical IV Analysis":
                self.show_historical_iv(hist_data, ticker)
            elif analysis_type == "Options Chain Analysis":
                self.show_options_chain(stock, ticker)
            elif analysis_type == "Strategy Analysis":
                self.show_strategy_analysis(stock, hist_data, ticker)
                
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def show_historical_iv(self, hist_data, ticker):
        # Display summary statistics
        summary = ttk.LabelFrame(self.analysis_frame, text="IV Summary")
        summary.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
        
        current_iv = hist_data['Historical_IV'].iloc[-1]
        avg_iv = hist_data['Historical_IV'].mean()
        max_iv = hist_data['Historical_IV'].max()
        min_iv = hist_data['Historical_IV'].min()
        
        ttk.Label(summary, 
                 text=f"Current IV: {current_iv:.2f}%").grid(row=0, column=0)
        ttk.Label(summary, 
                 text=f"Average IV: {avg_iv:.2f}%").grid(row=1, column=0)
        ttk.Label(summary, 
                 text=f"Max IV: {max_iv:.2f}%").grid(row=2, column=0)
        ttk.Label(summary, 
                 text=f"Min IV: {min_iv:.2f}%").grid(row=3, column=0)
        
        # Plot IV
        self.plot_iv(hist_data, ticker)

    def show_options_chain(self, stock, ticker):
        # Get options expiration dates
        expirations = stock.options
        
        # Create expiration date dropdown
        ttk.Label(self.analysis_frame, 
                 text="Select Expiration:").grid(row=0, column=0)
        exp_var = tk.StringVar()
        exp_combo = ttk.Combobox(self.analysis_frame, 
                                textvariable=exp_var, 
                                values=expirations)
        exp_combo.grid(row=0, column=1)
        exp_combo.set(expirations[0])
        
        def update_chain(*args):
            chain = stock.option_chain(exp_var.get())
            
            # Display calls
            calls_frame = ttk.LabelFrame(self.analysis_frame, text="Calls")
            calls_frame.grid(row=1, column=0, sticky="nsew", padx=5, pady=5)
            
            calls = chain.calls[['strike', 'lastPrice', 'impliedVolatility', 
                               'volume', 'openInterest']]
            calls['impliedVolatility'] = calls['impliedVolatility'] * 100
            
            tree = ttk.Treeview(calls_frame)
            tree["columns"] = list(calls.columns)
            tree["show"] = "headings"
            
            for column in calls.columns:
                tree.heading(column, text=column)
                tree.column(column, width=100)
            
            for idx, row in calls.iterrows():
                tree.insert("", "end", values=list(row.round(2)))
            
            # Add scrollbar
            scrollbar = ttk.Scrollbar(calls_frame, orient="vertical", 
                                    command=tree.yview)
            scrollbar.pack(side="right", fill="y")
            tree.configure(yscrollcommand=scrollbar.set)
            tree.pack(fill="both", expand=True)
            
        exp_combo.bind('<<ComboboxSelected>>', update_chain)
        update_chain()

    def plot_iv(self, hist_data, ticker):
        # Clear previous chart
        for widget in self.chart_frame.winfo_children():
            widget.destroy()
            
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.plot(hist_data.index, hist_data['Historical_IV'], 
                label='Historical IV')
        ax.plot(hist_data.index, 
                hist_data['Historical_IV'].rolling(window=20).mean(),
                label='20-day MA', linestyle='--')
        
        ax.set_title(f'{ticker} Historical Implied Volatility')
        ax.set_xlabel('Date')
        ax.set_ylabel('IV (%)')
        ax.legend()
        ax.grid(True)
        
        canvas = FigureCanvasTkAgg(fig, self.chart_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    def export_data(self):
        ticker = self.ticker_entry.get().strip().upper()
        if not ticker:
            messagebox.showerror("Error", "Please run analysis first")
            return
            
        try:
            stock = yf.Ticker(ticker)
            hist_data = stock.history(period='1y')
            
            # Export to Excel
            filename = f"{ticker}_options_analysis.xlsx"
            with pd.ExcelWriter(filename) as writer:
                hist_data.to_excel(writer, sheet_name='Historical Data')
                
                # Add options data if available
                if stock.options:
                    chain = stock.option_chain(stock.options[0])
                    chain.calls.to_excel(writer, sheet_name='Calls')
                    chain.puts.to_excel(writer, sheet_name='Puts')
                    
            messagebox.showinfo("Success", f"Data exported to {filename}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Export failed: {str(e)}")

    def export_report(self):
        # Create Word document report
        doc = Document()
        doc.add_heading('Options Analysis Report', 0)
        
        ticker = self.ticker_entry.get().strip().upper()
        doc.add_paragraph(f"Analysis for {ticker}")
        
        # Add analysis details based on type
        analysis_type = self.analysis_type.get()
        doc.add_heading(analysis_type, level=1)
        
        # Save report
        filename = f"{ticker}_options_report.docx"
        doc.save(filename)
        messagebox.showinfo("Success", f"Report exported to {filename}")

    def run(self):
        self.root.mainloop()

if __name__ == "__main__":
    app = OptionsAnalyzer()
    app.run()